"""
Engine3: DocxTranslatorEngine
==============================
Menerjemahkan dokumen Word dari bahasa asing ke Bahasa Indonesia.

PRINSIP UTAMA:
  - HANYA menerjemahkan teks saja.
  - TIDAK mengubah apapun: margin, font, ukuran, bold, italic, spacing,
    alignment, style, heading, tabel, gambar, header, footer, section, dll.
  - Struktur dan tampilan dokumen output IDENTIK dengan dokumen asli.
  - Teks yang tidak perlu diterjemahkan (angka, kode, copyright, dll)
    dibiarkan apa adanya.

Dependensi:
    pip install deep-translator python-docx
"""

import re
import time
import traceback
from docx import Document
from deep_translator import GoogleTranslator


# ─────────────────────────────────────────────────────────────────────────────
# KONSTANTA & POLA REGEX
# ─────────────────────────────────────────────────────────────────────────────

# Teks yang TIDAK diterjemahkan
_RE_PURE_NUMBER  = re.compile(r'^[\d\s\.\,\:\;\-\(\)\[\]\/\\\+\=\*\%\&\^\$\#\@\!\"\'`~<>{}|_]+$')
_RE_COPYRIGHT    = re.compile(r'©|All\s+rights\s+reserved', re.IGNORECASE)
_RE_SNI_CODE     = re.compile(r'^(SNI|ISO|IEC|ASTM|BS|DIN|JIS|EN|ANSI)\s', re.IGNORECASE)

# Style paragraf yang TIDAK diterjemahkan
_SKIP_STYLES = {
    'toc 1', 'toc 2', 'toc 3', 'toc 4', 'toc 5',
    'table of figures',
    'footnote text', 'endnote text',
    'macro text',
}

# Namespace Word
_NS_W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

# Jeda antar request ke Google Translate (detik)
_TRANSLATE_DELAY = 0.15


# ─────────────────────────────────────────────────────────────────────────────
# HELPER: FILTER TEKS
# ─────────────────────────────────────────────────────────────────────────────

def _skip_text(text: str) -> bool:
    """Return True jika teks tidak perlu diterjemahkan."""
    t = text.strip()
    if len(t) < 3:
        return True
    if _RE_PURE_NUMBER.fullmatch(t):
        return True
    if _RE_COPYRIGHT.search(t):
        return True
    if _RE_SNI_CODE.match(t):
        return True
    return False


def _skip_paragraph(para) -> bool:
    """Return True jika paragraf dilewati seluruhnya."""
    # Kosong
    if not para.text.strip():
        return True
    # Mengandung gambar/drawing
    for tag in [f'{{{_NS_W}}}drawing', f'{{{_NS_W}}}pict']:
        if para._element.find('.//' + tag) is not None:
            return True
    # Style khusus
    style_name = (para.style.name or '').lower()
    if any(style_name.startswith(s) for s in _SKIP_STYLES):
        return True
    return False


# ─────────────────────────────────────────────────────────────────────────────
# HELPER: DETEKSI ITALIC (untuk proteksi cover)
# ─────────────────────────────────────────────────────────────────────────────

def _all_runs_italic(para) -> bool:
    """
    Return True jika semua run yang memiliki teks bersifat italic.
    Digunakan untuk melindungi judul bahasa asing (italic) di halaman cover.
    """
    text_runs = [r for r in para.runs if r.text.strip()]
    if not text_runs:
        return False

    para_style_italic = False
    try:
        if para.style and para.style.font and para.style.font.italic:
            para_style_italic = True
    except Exception:
        pass

    for run in text_runs:
        run_italic = False

        # Cek via API python-docx
        if run.font.italic is True:
            run_italic = True

        # Cek langsung via XML <w:i/>
        if not run_italic:
            rPr = run._element.find(f'{{{_NS_W}}}rPr')
            if rPr is not None:
                i_el = rPr.find(f'{{{_NS_W}}}i')
                if i_el is not None:
                    val = i_el.get(f'{{{_NS_W}}}val', 'true')
                    if val.lower() not in ('false', '0'):
                        run_italic = True

        # Fallback warisan dari style paragraf
        if not run_italic and para_style_italic:
            rPr = run._element.find(f'{{{_NS_W}}}rPr')
            if rPr is not None:
                i_el = rPr.find(f'{{{_NS_W}}}i')
                if i_el is not None:
                    val = i_el.get(f'{{{_NS_W}}}val', 'true')
                    if val.lower() not in ('false', '0'):
                        run_italic = True
                else:
                    run_italic = True
            else:
                run_italic = True

        if not run_italic:
            return False

    return True


def _has_inline_sectpr(para) -> bool:
    """Return True jika paragraf ini menandai akhir section (batas cover)."""
    pPr = para._element.find(f'{{{_NS_W}}}pPr')
    if pPr is None:
        return False
    return pPr.find(f'{{{_NS_W}}}sectPr') is not None


# ─────────────────────────────────────────────────────────────────────────────
# CORE: TERJEMAHKAN PARAGRAF (run by run, format tetap)
# ─────────────────────────────────────────────────────────────────────────────

def _translate_paragraph(para, translator: GoogleTranslator) -> None:
    """
    Terjemahkan teks dalam paragraf sambil mempertahankan format setiap run.

    Strategi:
      1. Gabungkan teks dari semua run yang punya teks.
      2. Terjemahkan hasil gabungan sebagai satu unit.
      3. Letakkan hasil terjemahan di run pertama yang berteks.
      4. Kosongkan run-run lainnya — format (bold, italic, font, dll) TETAP ada.
    """
    if _skip_paragraph(para):
        return

    text_runs = [(i, r) for i, r in enumerate(para.runs) if r.text]
    if not text_runs:
        return

    combined = ''.join(r.text for _, r in text_runs)
    if _skip_text(combined):
        return

    try:
        translated = translator.translate(combined.strip())
        time.sleep(_TRANSLATE_DELAY)
    except Exception:
        time.sleep(1.0)
        try:
            translated = translator.translate(combined.strip())
        except Exception:
            return  # gagal dua kali → biarkan teks asli

    if not translated or translated == combined:
        return

    # Letakkan hasil terjemahan di run pertama
    _, first_run = text_runs[0]
    first_run.text = translated

    # Kosongkan run sisanya (format tetap, hanya teks dihapus)
    for _, run in text_runs[1:]:
        run.text = ''


# ─────────────────────────────────────────────────────────────────────────────
# CORE: TERJEMAHKAN TABEL
# ─────────────────────────────────────────────────────────────────────────────

def _translate_table(table, translator: GoogleTranslator) -> None:
    """Terjemahkan teks dalam seluruh sel tabel."""
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                _translate_paragraph(para, translator)


# ─────────────────────────────────────────────────────────────────────────────
# MAIN ENGINE CLASS
# ─────────────────────────────────────────────────────────────────────────────

class DocxTranslatorEngine:
    """
    Engine 3 — Terjemahkan dokumen Word ke Bahasa Indonesia.

    HANYA menerjemahkan teks. Tidak ada perubahan format apapun.
    Dokumen output identik strukturnya dengan dokumen asli.

    Cara pakai:
        engine3 = DocxTranslatorEngine()
        success, msg = engine3.process(input_path, output_path)
    """

    def __init__(self, source_lang: str = 'auto', target_lang: str = 'id'):
        self.source_lang = source_lang
        self.target_lang = target_lang

    def process(
        self,
        input_path:  str,
        output_path: str,
        font_name:   str = None,   # diabaikan, hanya untuk kompatibilitas signature lama
        font_size:   int = None,   # diabaikan, hanya untuk kompatibilitas signature lama
    ) -> tuple[bool, str]:
        """
        Terjemahkan dokumen dan simpan hasilnya.

        Args:
            input_path  : path file .docx sumber
            output_path : path file .docx output
            font_name   : tidak digunakan (tetap diterima agar kompatibel dengan app.py)
            font_size   : tidak digunakan (tetap diterima agar kompatibel dengan app.py)

        Returns:
            (True, output_path) jika berhasil
            (False, pesan_error) jika gagal
        """
        try:
            print("Engine3: Memulai penerjemahan (format dokumen tidak akan diubah)...")

            # ── Init GoogleTranslator ──────────────────────────────────────
            translator = GoogleTranslator(source=self.source_lang, target=self.target_lang)

            # ── Buka dokumen ───────────────────────────────────────────────
            doc = Document(input_path)

            # ── Deteksi batas cover section ───────────────────────────────
            # Paragraf italic di section cover (judul bahasa asing) tidak diterjemahkan.
            # Cover dari engine4 selalu ditutup oleh paragraf dengan inline <w:sectPr>.
            section_breaks_seen = 0
            COVER_SECTION_END   = 1  # batas section break yang menandai akhir cover

            # ── Iterasi body: paragraf & tabel ─────────────────────────────
            body = doc.element.body
            para_map  = {p._element: p for p in doc.paragraphs}
            table_map = {t._element: t for t in doc.tables}

            total = sum(1 for c in body if c in para_map or c in table_map)
            done  = 0

            for child in body:

                if child in para_map:
                    para     = para_map[child]
                    in_cover = (section_breaks_seen < COVER_SECTION_END)

                    # Judul bahasa asing (italic) di cover → LEWATI
                    if in_cover and _all_runs_italic(para):
                        print(f"  [Cover-italic dilewati]: {para.text[:60]}")
                    else:
                        _translate_paragraph(para, translator)

                    # Tandai section break
                    if _has_inline_sectpr(para):
                        section_breaks_seen += 1

                    done += 1
                    if done % 20 == 0:
                        print(f"  Progress: {done}/{total} elemen diproses...")

                elif child in table_map:
                    _translate_table(table_map[child], translator)
                    done += 1

            # ── Simpan — TIDAK ada perubahan format sama sekali ────────────
            doc.save(output_path)
            print(f"Engine3: Selesai. Output disimpan ke → {output_path}")
            return True, output_path

        except ImportError:
            return False, (
                "deep-translator tidak terinstall.\n"
                "Jalankan: pip install deep-translator"
            )
        except Exception as e:
            return False, f"Engine3 Error: {str(e)}\n{traceback.format_exc()}"